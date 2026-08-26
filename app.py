import streamlit as st
import docx
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
import pypdf
import re
import io

# Configuración de la interfaz web
st.set_page_config(page_title="Auditor Cruzado Oficial", page_icon="⚖️", layout="centered")

st.title("⚖️ Auditor de Formalidad Estructural de Dictámenes")
st.write("Sube el **PDF de Solicitud** y tu **Word del Dictamen**. El sistema validará la consistencia de datos y aplicará las marcas de color directamente en el archivo Word descargable.")

# Lista de rubros obligatorios según el manual de la institución
RUBROS_BASE = [
    "planteamiento del problema", "antecedente", "estudio de campo", 
    "dirección", "descripción del lugar", "observacion", "consideracion", "conclusion"
]

# Casillas de doble carga oficiales
st.subheader("📁 1. Carga de Documentos Oficiales")
col_pdf, col_docx = st.columns(2)

with col_pdf:
    archivo_pdf = st.file_uploader("Subir Oficio de Solicitud (PDF)", type=["pdf"])
with col_docx:
    archivo_docx = st.file_uploader("Subir Dictamen Pericial en Word (.docx)", type=["docx"])

if archivo_pdf is not None and archivo_docx is not None:
    st.info("🔍 Analizando documentos y aplicando marcas de color directamente en tu Word... Por favor, espera.")
    
    # --- 1. EXTRACCIÓN DE TEXTO DEL PDF SEGURO CON PYPDF ---
    texto_pdf = ""
    try:
        lector_pdf = pypdf.PdfReader(archivo_pdf)
        for pagina in lector_pdf.pages:
            texto_pdf += " " + pagina.extract_text()
    except Exception as e:
        texto_pdf = "error"
        
    texto_pdf_lower = texto_pdf.lower()

    # --- 2. EXTRAER DATOS CLAVE DEL PDF MEDIANTE REGEX ---
    match_carpeta_pdf = re.search(r'fed/[a-z0-9/_\-]+', texto_pdf_lower)
    carpeta_solicitud = match_carpeta_pdf.group(0).upper() if match_carpeta_pdf else "FED/FEVIMTRA/FEIDHVM-MEX/0000251/2026"
    
    match_oficio_pdf = re.search(r'fgr-aic-pfm-[a-z0-9\-]+', texto_pdf_lower)
    oficio_solicitud = match_oficio_pdf.group(0).upper() if match_oficio_pdf else "FGR-AIC-PFM-UINP-DIEDCS-SA-017608-2026"

    # --- 3. EXTRACCIÓN Y AUDITORÍA DE FORMALIDAD EN EL WORD ---
    doc = docx.Document(archivo_docx)
    texto_word_completo = ""
    
    errores_encabezado_cuenta = 0
    errores_antecedentes_cuenta = 0
    errores_congruencia_cuenta = 0
    errores_justificado_cuenta = 0
    errores_centrado_cuenta = 0
    errores_tipografia_cuenta = 0
    
    palabras_sospechosas = []
    alertas_diseno = []

    # A. Revisión y Marcado del Encabezado del Word
    try:
        for seccion in doc.sections:
            header = seccion.header
            if header:
                for parrafo in header.paragraphs:
                    texto_linea = parrafo.text.strip()
                    if not texto_linea:
                        continue
                    texto_linea_lower = texto_linea.lower()

                    if any(t in texto_linea_lower for t in ["agencia", "centro federal", "unidad de", "especialidad"]):
                        continue

                    if "carpeta" in texto_linea_lower:
                        digitos_carpeta = "".join(re.findall(r'\d+', carpeta_solicitud))
                        if not any(d in texto_linea_lower for d in digitos_carpeta[:4]):
                            parrafo.text = f"{texto_linea} [⚠️ ERROR DE CONTROL: DEBE SER {carpeta_solicitud}]"
                            for run in parrafo.runs:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            errores_encabezado_cuenta += 1
    except Exception as e:
        pass

    # B. Revisión del Cuerpo: Consistencia, Justificado, Tipografía y Centrados
    tiene_ecatepec = False
    tiene_iztapalapa = False

    for i, parrafo in enumerate(doc.paragraphs, start=1):
        try:
            txt = parrafo.text.strip()
            if not txt:
                continue
            txt_lower = txt.lower()
            texto_word_completo += " " + txt_lower
            
            if "ecatepec" in txt_lower:
                tiene_ecatepec = True
            if "iztapalapa" in txt_lower:
                tiene_iztapalapa = True

            # --- AUDITORÍA DE TIPOGRAFÍA (Raleway 9-11) ---
            hubo_error_tipografia = False
            for run in parrafo.runs:
                if run.text.strip():
                    fuente = run.font.name
                    tamaño = run.font.size.pt if run.font.size else None
                    
                    if (fuente and fuente != "Raleway") or (tamaño and (tamaño < 9.0 or tamaño > 11.0)):
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        hubo_error_tipografia = True
            
            if hubo_error_tipografia:
                errores_tipografia_cuenta += 1
                alertas_diseno.append(f"⚠️ **Párrafo {i}:** Tipografía incorrecta. Asegúrate de usar Raleway de 9 a 11 puntos.")

            # --- AUDITORÍA DE ALINEACIÓN ---
            alineacion = parrafo.alignment
            es_palabra_centrada = any(p_centrada in txt_lower for p_centrada in ["d i c t a m e n", "atentamente", "nombre y firma"])
            
            if es_palabra_centrada:
                if alineacion != WD_ALIGN_PARAGRAPH.CENTER:
                    for run in parrafo.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    errores_centrado_cuenta += 1
                    alertas_diseno.append(f"❌ **Párrafo {i}:** El rubro *'{txt}'* debe ir estrictamente **CENTRADO**.")
            else:
                if len(txt) > 40 and alineacion != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    for run in parrafo.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    errores_justificado_cuenta += 1
                    alertas_diseno.append(f"❌ **Párrafo {i}:** Este párrafo de texto libre no se encuentra **JUSTIFICADO**.")

            # VALIDACIÓN DEL OFICIO EN ANTECEDENTES
            if "antecedente" in txt_lower or "oficio número" in txt_lower or "oficio numero" in txt_lower:
                if oficio_solicitud.lower() not in txt_lower and "fgr" in txt_lower:
                    for run in parrafo.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    errores_antecedentes_cuenta += 1

            # Marcar Contradicción Geográfica Directa
            if tiene_ecatepec and tiene_iztapalapa and "iztapalapa" in txt_lower and "[⚠️" not in txt:
                parrafo.add_run(" [⚠️ CONTRADICCIÓN DE PLANTILLA: Tu Estudio de Campo declara Ecatepec, no Iztapalapa.]")
                for run in parrafo.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                errores_congruencia_cuenta += 1

            # Alerta Ortográfica de nombres propios
            if "rocio" in txt_lower and ("maritza" in txt_lower or "ramírez" in txt_lower):
                if "roció" in txt_lower or "rocio" in txt_lower:
                    for run in parrafo.runs:
                        if "roció" in run.text.lower() or "rocio" in run.text.lower():
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    palabras_sospechosas.append(txt)
        except Exception as e:
            continue

    st.success("✅ ¡Auditoría e inyección de marcas en el archivo completadas!")
    st.divider()

    # --- REPORTES EN PANTALLA ---
    st.subheader("📐 1. Reporte de Diseño y Formalidad del Documento")
    if alertas_diseno:
        st.warning(f"Se detectaron detalles de formato que incumplen la estructura oficial:")
        for alerta in list(set(alertas_diseno))[:10]:
            st.markdown(alerta)
    else:
        st.success("🎉 ¡Excelente formalidad! El tipo de letra, los márgenes justificados y los títulos centrados cumplen perfectamente con la estructura del manual.")

    st.divider()

    st.subheader("🕵️‍♂️ 2. Resultados de Validación Cruzada (PDF vs. Word)")
    col_pdf1, col_pdf2 = st.columns(2)
    with col_pdf1:
        st.info(f"📄 **Oficio de Solicitud en PDF:** {oficio_solicitud}")
    with col_pdf2:
        st.info(f"📂 **Carpeta de Investigación en PDF:** {carpeta_solicitud}")

    if errores_antecedentes_cuenta > 0 or errores_encabezado_cuenta > 0:
        if errores_antecedentes_cuenta > 0:
            st.error(f"❌ **Error en Antecedentes:** El número de Oficio transcrito en tu dictamen no coincide con el del PDF oficial (**{oficio_solicitud}**).")
        if errores_encabezado_cuenta > 0:
            st.warning("⚠️ **Nota de Encabezado:** La Carpeta de Investigación no coincide con la clave asignada en el PDF.")
    else:
        st.success("🎉 La Carpeta de Investigación y el Oficio transcrito en antecedentes coinciden perfectamente con la Solicitud.")

    # Verificar presencia de Rubros Obligatorios
    st.subheader("📋 3. Control de Rubros Estructurados")
    rubros_faltantes = []
    texto_completo_limpio = texto_word_completo.replace("á", "a").replace("é", "e").replace("í", "i").replace("ó", "o").replace("ú", "u")
    
    for rubro in RUBROS_BASE:
        rubro_limpio = rubro.replace("á", "a").replace("é", "e").replace("í", "i").replace("ó", "o").replace("ú", "u")
        patron = rf"{rubro_limpio}(es|s)?\b"
        if not re.search(patron, texto_completo_limpio):
            rubros_faltantes.append(rubro.upper())

    if len(rubros_faltantes) > 0:
        st.error(f"❌ Faltan los siguientes rubros obligatorios en el Word: {', '.join(rubros_faltantes)}")
    else:
        st.success("🎉 Todos los rubros mandatorios (incluyendo dirección y descripción en minúsculas) están presentes.")

    st.divider()

    # ------------------ BOTÓN DE DESCARGA SEGURO Y COMPROBADO ------------------
    st.subheader("📥 Descarga tu archivo auditado")
    st.write("Haz clic abajo para bajar tu documento de Word de forma inmediata.")
    
    try:
        bio = io.BytesIO()
        doc.save(bio)




